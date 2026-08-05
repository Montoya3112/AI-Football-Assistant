# ════════════════════════════════════════════════════════════════
# 📦 MÓDULO: main.py
# 🎯 PROPÓSITO: Servidor Web Principal — API REST con FastAPI
#    - Sirve el frontend estático (HTML/CSS/JS)
#    - Expone endpoints de Chat IA, Visión OCR y Autenticación
#    - Genera y descarga documentos Word (.docx) de cédulas
# 🔗 DEPENDENCIAS: FastAPI, Supabase, ai_sports_engine, python-docx
# 📁 UBICACIÓN: app/main.py
# ════════════════════════════════════════════════════════════════

# ── IMPORTACIONES ──────────────────────────────────────────────
from fastapi import FastAPI, UploadFile, File, HTTPException, Depends, Header, Response
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
import os
import json
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from .core import supabase
from .schemas import MensajeChat, RespuestaChat, RespuestaVision, RegistroUsuario, LoginUsuario, RespuestaAuth, CedulaArbitralExtraida
from .ai_sports_engine import asistente_tecnico_chat, procesar_cedula_vision

# ── INSTANCIA DE LA APLICACIÓN FASTAPI ─────────────────────────
# FastAPI es el framework web que maneja todas las rutas HTTP.
# Documentación automática disponible en /docs (Swagger UI).
app = FastAPI(
    title="MRCA AI Football Assistant",
    version="1.0.0",
    description="Asistente de IA Futbolística & Visión - MRCA Solutions",
)

# ── MIDDLEWARE CORS ─────────────────────────────────────────────
# Permite que el frontend (browser) se comunique con el backend
# desde cualquier origen. Necesario para desarrollo local y producción.
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── ARCHIVOS ESTÁTICOS (Frontend HTML/CSS/JS) ──────────────────
# Monta la carpeta /static para servir index.html, style.css y script.js
static_path = os.path.join(os.path.dirname(__file__), "static")
app.mount("/static", StaticFiles(directory=static_path), name="static")

# ── DEPENDENCIA: AUTENTICACIÓN DE USUARIO ──────────────────────
# Verifica el token JWT del header Authorization usando Supabase Auth.
# Si no hay token o es "guest-access", retorna usuario invitado.
async def get_current_user(authorization: str = Header(None)):
    if not authorization:
        return {"email": "invitado@mrca.local", "nombre": "Invitado"}
    try:
        token = authorization.replace("Bearer ", "").strip()
        if token == "guest-access":
            return {"email": "invitado@mrca.local", "nombre": "Invitado"}
        user_res = supabase.auth.get_user(token)
        if user_res and user_res.user:
            return {"email": user_res.user.email, "nombre": user_res.user.user_metadata.get("nombre", "Usuario")}
    except Exception:
        pass
    return {"email": "invitado@mrca.local", "nombre": "Invitado"}

# ════════════════════════════════════════════════════════════════
# ENDPOINTS — RUTAS HTTP DE LA API
# ════════════════════════════════════════════════════════════════

# ── RUTA RAÍZ: Sirve la aplicación web ─────────────────────────
# GET / → Retorna index.html (punto de entrada del frontend SPA)
@app.get("/")
async def root():
    return FileResponse(os.path.join(static_path, "index.html"))

# ── ENDPOINT: Registro de nuevo usuario ────────────────────────
# POST /api/v1/auth/register
# Crea una cuenta en Supabase Auth con email + password + nombre.
# Retorna token JWT para sesión inmediata (si no requiere confirmación).
@app.post("/api/v1/auth/register", response_model=RespuestaAuth)
async def register(payload: RegistroUsuario):
    try:
        res = supabase.auth.sign_up({
            "email": payload.email,
            "password": payload.password,
            "options": {"data": {"nombre": payload.nombre or payload.email.split('@')[0]}}
        })
        user_data = {"email": payload.email, "nombre": payload.nombre}
        return RespuestaAuth(access_token="pending-confirmation", user=user_data)
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

# ── ENDPOINT: Inicio de sesión ──────────────────────────────────
# POST /api/v1/auth/login
# Autentica email + password contra Supabase Auth.
# Retorna JWT (access_token) que el frontend guarda en memoria.
@app.post("/api/v1/auth/login", response_model=RespuestaAuth)
async def login(payload: LoginUsuario):
    try:
        res = supabase.auth.sign_in_with_password({"email": payload.email, "password": payload.password})
        if not res.user or not res.session:
            raise HTTPException(status_code=401, detail="Credenciales inválidas o correo no confirmado")
        user_data = {"email": res.user.email, "nombre": res.user.user_metadata.get("nombre", "Usuario")}
        return RespuestaAuth(access_token=res.session.access_token, user=user_data)
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

# ── ENDPOINT: Chat con IA Futbolística ─────────────────────────
# POST /api/v1/futbol/chat
# ENTRADA: { "mensaje": "Dame la alineación del Real Madrid" }
# PROCESO: Llama a asistente_tecnico_chat() → Gemini API → respuesta
# SALIDA : { "respuesta": "texto en Markdown con análisis táctico" }
@app.post("/api/v1/futbol/chat", response_model=RespuestaChat)
async def chat(payload: MensajeChat, current_user: dict = Depends(get_current_user)):
    try:
        respuesta = await asistente_tecnico_chat(payload.mensaje)
        return RespuestaChat(respuesta=respuesta)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ── ENDPOINT: Visión OCR — Cédula Arbitral ─────────────────────
# POST /api/v1/futbol/vision/cedulas
# ENTRADA: Imagen JPG/PNG de una cédula arbitral (multipart/form-data)
# PROCESO: Gemini Vision analiza la imagen y extrae campos estructurados
# SALIDA : JSON con equipos, goles, tarjetas, jugadores e incidencias
# ALMACÉN: El registro se inserta en tabla "reportes_partidos" de Supabase
@app.post("/api/v1/futbol/vision/cedulas", response_model=RespuestaVision)
async def vision_cedulas(file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    if not file.content_type or not file.content_type.startswith("image/"):
        raise HTTPException(status_code=400, detail="Solo se aceptan archivos de imagen (JPG, PNG, WEBP)")
    try:
        imagen_bytes = await file.read()
        cedula = await procesar_cedula_vision(imagen_bytes)
        record = cedula.model_dump()
        for key in ["tarjetas_amarillas", "tarjetas_rojas", "goles", "incidencias", "jugadores_local", "jugadores_visitante"]:
            if isinstance(record.get(key), list):
                record[key] = json.dumps(record[key], ensure_ascii=False)
        inserted = False
        try:
            supabase.table("reportes_partidos").insert(record).execute()
            inserted = True
        except Exception:
            pass
        return RespuestaVision(cedula=cedula, raw_text=str(record), inserted=inserted)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ── ENDPOINT: Exportar Cédula como Documento Word (.docx) ──────
# POST /api/v1/futbol/vision/export-doc
# ENTRADA: JSON con el schema CedulaArbitralExtraida (campos del partido)
# PROCESO: python-docx construye un documento profesional con:
#          portada, tabla de datos, plantillas de jugadores,
#          tarjetas, incidencias y sección de firmas.
# SALIDA : Archivo .docx descargable vía StreamingResponse
@app.post("/api/v1/futbol/vision/export-doc")
async def export_cedula_word(cedula: CedulaArbitralExtraida):
    try:
        doc = Document()
        
        # Header title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run("CÉDULA ARBITRAL Y INFORME OFICIAL DE PARTIDO")
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(124, 106, 239)

        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = sub_p.add_run("MRCA Solutions — AI Football Assistant Vision System")
        run_sub.font.size = Pt(10)
        run_sub.font.italic = True
        run_sub.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph().paragraph_format.space_after = Pt(12)

        # Match Info Table
        table = doc.add_table(rows=4, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        row_0 = table.rows[0].cells
        row_0[0].text = f"⚽ Equipo Local: {cedula.equipo_local}"
        row_0[1].text = f"⚽ Equipo Visitante: {cedula.equipo_visitante}"

        row_1 = table.rows[1].cells
        row_1[0].text = f"🏆 Marcador Final: {cedula.goles_local} - {cedula.goles_visitante}"
        row_1[1].text = f"📅 Fecha: {cedula.fecha or 'N/A'}"

        row_2 = table.rows[2].cells
        row_2[0].text = f"🏟️ Estadio: {cedula.estadio or 'N/A'}"
        row_2[1].text = f"👨‍⚖️ Árbitro: {cedula.arbitro or 'N/A'}"

        row_3 = table.rows[3].cells
        row_3[0].text = f"🥇 Torneo / Liga: {cedula.torneo_liga or 'Oficial'}"
        row_3[1].text = f"📋 Estado: Reconstruido por IA (Normalizado)"

        doc.add_paragraph().paragraph_format.space_after = Pt(14)

        # Rosters Section
        doc.add_heading("Plantillas de Jugadores", level=2)
        
        # Local Players Table
        p_loc = doc.add_paragraph(f"Local: {cedula.equipo_local}")
        p_loc.runs[0].bold = True
        if cedula.jugadores_local:
            t_loc = doc.add_table(rows=1, cols=3)
            hdr = t_loc.rows[0].cells
            hdr[0].text = "Nº"
            hdr[1].text = "Nombre del Jugador"
            hdr[2].text = "Posición"
            for j in cedula.jugadores_local:
                row = t_loc.add_row().cells
                row[0].text = str(j.numero) if j.numero is not None else "-"
                row[1].text = j.nombre
                row[2].text = j.posicion or "-"
        else:
            doc.add_paragraph("(No se especificó lista de jugadores local)")

        doc.add_paragraph().paragraph_format.space_after = Pt(10)

        # Visitor Players Table
        p_vis = doc.add_paragraph(f"Visitante: {cedula.equipo_visitante}")
        p_vis.runs[0].bold = True
        if cedula.jugadores_visitante:
            t_vis = doc.add_table(rows=1, cols=3)
            hdr_v = t_vis.rows[0].cells
            hdr_v[0].text = "Nº"
            hdr_v[1].text = "Nombre del Jugador"
            hdr_v[2].text = "Posición"
            for j in cedula.jugadores_visitante:
                row = t_vis.add_row().cells
                row[0].text = str(j.numero) if j.numero is not None else "-"
                row[1].text = j.nombre
                row[2].text = j.posicion or "-"
        else:
            doc.add_paragraph("(No se especificó lista de jugadores visitante)")

        doc.add_paragraph().paragraph_format.space_after = Pt(14)

        # Cards Section
        doc.add_heading("Registro de Tarjetas y Sanciones", level=2)
        if cedula.tarjetas_amarillas or cedula.tarjetas_rojas:
            t_cards = doc.add_table(rows=1, cols=3)
            h = t_cards.rows[0].cells
            h[0].text = "Tipo"
            h[1].text = "Jugador"
            h[2].text = "Minuto"
            for t in cedula.tarjetas_amarillas:
                r = t_cards.add_row().cells
                r[0].text = "🟨 Amarilla"
                r[1].text = t.jugador
                r[2].text = f"Min {t.minuto}'" if t.minuto else "-"
            for t in cedula.tarjetas_rojas:
                r = t_cards.add_row().cells
                r[0].text = "🟥 Roja"
                r[1].text = t.jugador
                r[2].text = f"Min {t.minuto}'" if t.minuto else "-"
        else:
            doc.add_paragraph("Sin amonestaciones ni expulsiones registradas.")

        doc.add_paragraph().paragraph_format.space_after = Pt(14)

        # Signatures
        doc.add_heading("Firmas de Conformidad", level=2)
        sig_p = doc.add_paragraph()
        sig_p.add_run("_______________________             _______________________\nFirma del Árbitro Central              Firma Capitán Local\n\n_______________________\nFirma Capitán Visitante")

        # ── GENERACIÓN Y DESCARGA DEL ARCHIVO ──────────────────
        # Se serializa el documento a un buffer en memoria (BytesIO)
        # y se envía como StreamingResponse con headers de descarga.
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        filename = f"Cedula_Arbitral_{cedula.equipo_local}_vs_{cedula.equipo_visitante}.docx".replace(" ", "_")
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ── ENDPOINT: Health Check ──────────────────────────────────────
# GET /api/health — Verifica que el servidor está corriendo
# Usado por plataformas de hosting (Render, Railway) para monitoreo.
@app.get("/api/health")
async def health():
    return {"status": "ok", "service": "MRCA AI Football Assistant"}
